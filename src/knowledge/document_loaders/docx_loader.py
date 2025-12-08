# -*- coding: utf-8 -*-
"""
@File    : docx_loader.py
@Time    : 2025/12/8 16:02
@Desc    : 
"""
from typing import List, Optional
import zipfile
import xml.etree.ElementTree as ET
from .base_loader import BaseDocumentLoader, Document


class DocxLoader(BaseDocumentLoader):
    """Word文档(.docx)加载器"""

    def __init__(self, file_path: str, encoding: str = "utf-8",
                 extract_images: bool = False, extract_tables: bool = False):
        """
        初始化DOCX加载器

        Args:
            file_path: 文件路径
            encoding: 文件编码
            extract_images: 是否提取图片描述
            extract_tables: 是否提取表格内容
        """
        super().__init__(file_path, encoding)
        self.extract_images = extract_images
        self.extract_tables = extract_tables

    def load(self) -> List[Document]:
        """加载DOCX文档"""
        try:
            # 使用python-docx库（如果可用）
            try:
                import docx
                return self._load_with_docx()
            except ImportError:
                # 回退到自定义解析
                return self._load_with_xml()

        except Exception as e:
            raise Exception(f"加载DOCX文件失败: {str(e)}")

    def _load_with_docx(self) -> List[Document]:
        """使用python-docx库加载"""
        import docx

        # 打开文档
        doc = docx.Document(self.file_path)

        # 提取所有段落文本
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # 跳过空段落
                full_text.append(para.text)

        # 提取表格内容（如果启用）
        if self.extract_tables:
            table_texts = self._extract_tables_with_docx(doc)
            full_text.extend(table_texts)

        # 提取图片描述（如果启用）
        if self.extract_images:
            image_descriptions = self._extract_images_with_docx(doc)
            full_text.extend(image_descriptions)

        # 合并所有文本
        content = "\n".join(full_text)

        # 提取文档属性
        metadata = {
            "source": str(self.file_path),
            "file_type": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables) if hasattr(doc, 'tables') else 0
        }

        # 尝试提取标题
        title = self._extract_docx_title(doc)
        if title:
            metadata["title"] = title

        # 尝试提取作者等信息
        if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
            metadata["author"] = doc.core_properties.author
        if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
            metadata["created_date"] = doc.core_properties.created.isoformat()

        document = Document(
            content=content,
            metadata=metadata
        )

        return [document]

    def _extract_tables_with_docx(self, doc) -> List[str]:
        """使用python-docx提取表格内容"""
        table_texts = []

        if not hasattr(doc, 'tables'):
            return table_texts

        for table_idx, table in enumerate(doc.tables, 1):
            table_rows = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    table_rows.append(" | ".join(row_cells))

            if table_rows:
                table_text = f"表格 {table_idx}:\n" + "\n".join(table_rows)
                table_texts.append(table_text)

        return table_texts

    def _extract_images_with_docx(self, doc) -> List[str]:
        """提取图片描述"""
        image_descriptions = []

        # 注意：python-docx不直接提供图片内容提取
        # 这里只能提取图片引用
        # 对于实际图片内容提取，需要更复杂的处理
        image_descriptions.append("【文档包含图片，建议查看原文获取完整信息】")

        return image_descriptions

    def _extract_docx_title(self, doc) -> Optional[str]:
        """从DOCX文档中提取标题"""
        # 方法1：从第一个段落中提取（假设是标题）
        if doc.paragraphs:
            first_para = doc.paragraphs[0].text.strip()
            if first_para and len(first_para) < 100:  # 假设标题不会太长
                return first_para

        # 方法2：从样式判断
        for para in doc.paragraphs:
            if hasattr(para, 'style') and para.style.name.lower().startswith('heading'):
                return para.text.strip()

        return None

    def _load_with_xml(self) -> List[Document]:
        """使用XML解析DOCX（不使用外部库）"""
        # DOCX实际上是ZIP文件，包含XML
        try:
            with zipfile.ZipFile(self.file_path, 'r') as docx_zip:
                # 读取主文档内容
                if 'word/document.xml' in docx_zip.namelist():
                    document_xml = docx_zip.read('word/document.xml')

                    # 解析XML
                    root = ET.fromstring(document_xml)

                    # 定义命名空间
                    namespaces = {
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    }

                    # 提取所有文本
                    text_elements = root.findall('.//w:t', namespaces)
                    texts = [elem.text for elem in text_elements if elem.text]

                    content = "\n".join(texts)

                    # 提取表格
                    if self.extract_tables:
                        table_texts = self._extract_tables_from_xml(root, namespaces)
                        content += "\n\n" + "\n".join(table_texts)

                    metadata = {
                        "source": str(self.file_path),
                        "file_type": "docx",
                        "parser": "xml"
                    }

                    document = Document(
                        content=content,
                        metadata=metadata
                    )

                    return [document]
                else:
                    raise Exception("DOCX文件格式错误：找不到document.xml")

        except zipfile.BadZipFile:
            raise Exception("文件不是有效的DOCX格式")
        except Exception as e:
            raise Exception(f"解析DOCX XML失败: {str(e)}")

    def _extract_tables_from_xml(self, root, namespaces) -> List[str]:
        """从XML中提取表格内容"""
        table_texts = []

        # 查找所有表格
        tables = root.findall('.//w:tbl', namespaces)

        for table_idx, table in enumerate(tables, 1):
            table_rows_text = []

            # 查找所有行
            rows = table.findall('.//w:tr', namespaces)

            for row in rows:
                row_cells_text = []

                # 查找所有单元格
                cells = row.findall('.//w:tc', namespaces)

                for cell in cells:
                    # 提取单元格中的文本
                    cell_text_elements = cell.findall('.//w:t', namespaces)
                    cell_texts = [elem.text for elem in cell_text_elements if elem.text]

                    if cell_texts:
                        row_cells_text.append(" ".join(cell_texts))

                if row_cells_text:
                    table_rows_text.append(" | ".join(row_cells_text))

            if table_rows_text:
                table_text = f"表格 {table_idx}:\n" + "\n".join(table_rows_text)
                table_texts.append(table_text)

        return table_texts


if __name__ == '__main__':
    loader = DocxLoader(r"D:\Coding\MyData\大模型技术栈-实战与应用.docx")
    docs = loader.load()
    print(docs)
